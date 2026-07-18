import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Matches either **bold** or [text](url) - whichever comes first, left to right
TOKEN_PATTERN = re.compile(r"\*\*(.+?)\*\*|\[([^\]]+)\]\(([^)]+)\)")


def _add_hyperlink(paragraph, url, text):
    """python-docx has no built-in hyperlink support - this builds the
    required XML manually (relationship + styled run) so links are real,
    clickable hyperlinks in Word, not plain bracket text."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_runs(paragraph, text):
    """Splits text on **bold** and [text](url) markers, adding each as a
    properly formatted run (bold text or real hyperlink) in order."""
    pos = 0
    for match in TOKEN_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])

        bold_text = match.group(1)
        link_text, link_url = match.group(2), match.group(3)

        if bold_text is not None:
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif link_text is not None:
            _add_hyperlink(paragraph, link_url, link_text)

        pos = match.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(markdown_text, title="FMCG Deal Intelligence Newsletter"):
    """
    Converts the newsletter's markdown structure into a Word document.
    Handles: # / ## / ### headings, "- " bullet lines, **bold** text, and
    [text](url) markdown links (converted to real clickable hyperlinks).
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")

    for raw_line in lines:
        line = raw_line.rstrip()

        if not line.strip():
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line.strip()[2:])
        else:
            p = doc.add_paragraph()
            _add_runs(p, line.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
